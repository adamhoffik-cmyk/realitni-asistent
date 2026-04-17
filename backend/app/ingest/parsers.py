"""Parsery pro PDF, DOCX, TXT → plain text."""
from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_pdf(path: Path) -> str:
    """Extrahuje text z PDF přes pdfplumber (lepší formátování než pypdf)."""
    import pdfplumber  # lazy

    text_parts: list[str] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                try:
                    t = page.extract_text()
                except Exception as exc:  # noqa: BLE001
                    logger.warning("%s: page extract error: %s", path.name, exc)
                    continue
                if t:
                    text_parts.append(t)
    except Exception as exc:  # noqa: BLE001
        logger.error("PDF parse failed for %s: %s", path, exc)
        return ""
    return "\n\n".join(text_parts)


def parse_docx(path: Path) -> str:
    """Extrahuje text z DOCX (paragraphs + tables)."""
    from docx import Document  # lazy

    try:
        doc = Document(str(path))
    except Exception as exc:  # noqa: BLE001
        logger.error("DOCX parse failed for %s: %s", path, exc)
        return ""

    parts: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells)
            if row_text.strip("| "):
                parts.append(row_text)
    return "\n".join(parts)


def parse_txt(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        try:
            return path.read_text(encoding="cp1250")
        except Exception as exc:  # noqa: BLE001
            logger.error("TXT parse failed for %s: %s", path, exc)
            return ""


def parse_file(path: Path) -> str:
    """Autodetekce parseru podle extension."""
    ext = path.suffix.lower()
    if ext == ".pdf":
        return parse_pdf(path)
    if ext in {".docx"}:
        return parse_docx(path)
    if ext in {".txt", ".md"}:
        return parse_txt(path)
    # Přeskoč ostatní (.doc je binary, .odt taky, images)
    return ""


# --------------------------------------------------------------------------- #
# Chunking — rozdělení dlouhého textu na kusy ~800 znaků s overlappem 100
# --------------------------------------------------------------------------- #
def chunk_text(
    text: str,
    chunk_size: int = 800,
    overlap: int = 100,
) -> list[str]:
    """Rozdělí text na chunks s překryvem.

    Respektuje hranice odstavců, pokud to jde.
    """
    text = text.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    paragraphs = text.split("\n\n")
    buf = ""
    for para in paragraphs:
        if len(buf) + len(para) + 2 > chunk_size and buf:
            chunks.append(buf.strip())
            # Overlap: vezmi posledních `overlap` znaků z předchozího chunku
            buf = buf[-overlap:] + "\n\n" + para
        else:
            buf = (buf + "\n\n" + para) if buf else para
    if buf.strip():
        chunks.append(buf.strip())

    return chunks
